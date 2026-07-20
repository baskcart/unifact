"""Generate UniFact training facilitator guide and role lab workbooks (.docx)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

OUT_DIR = Path(__file__).resolve().parents[1] / "docs" / "training"


def add_page_numbers(doc: Document) -> None:
    """Centered PAGE field in footer for every section."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(10)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(end)


def set_doc_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.footer_distance = Inches(0.5)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    add_page_numbers(doc)


def cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Arial"
    p2 = doc.add_paragraph(subtitle)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag = doc.add_paragraph("UniFact — One Fact. One Truth.")
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(11)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        para = doc.add_paragraph(item, style="List Bullet")
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, hdr in enumerate(headers):
        t.rows[0].cells[i].text = hdr
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def exercise(doc: Document, num: str, title: str, steps: list[str], record_lines: int = 0) -> None:
    h3(doc, f"Exercise {num}: {title}")
    bullets(doc, steps)
    if record_lines:
        p(doc, "Record your results below:", bold=True)
        for i in range(record_lines):
            doc.add_paragraph("_" * 72)


def build_foundation_handout() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    cover(doc, "UniFact Training", "Foundation Handout — All Roles")

    h1(doc, "Shared vocabulary")
    table(
        doc,
        ["Term", "Definition"],
        [
            ["Registry", "Tenancy — who owns facts and who may write"],
            ["Namespace", "Topic folder inside a registry (e.g. sales.policy)"],
            ["Uni Fact", "One named entry: namespace/key + value + metadata"],
            ["Published", "Approved truth — safe for prod, agents, and customers"],
            ["Proposed", "Draft awaiting review — not organizational truth"],
            ["Lookup", "Read-only link to another registry's published facts"],
            ["Membership", "Required to write/push; lookup alone is not enough"],
            ["Fact Check", "sync_pull → search/list facts before acting"],
        ],
    )

    h1(doc, "Fact lifecycle")
    p(doc, "working / proposed → review → published → supersede or retract")
    p(doc, "Rule: Never treat proposed or working facts as published truth.")

    h1(doc, "Resolution order")
    bullets(
        doc,
        [
            "Exact namespace in your registry",
            "Parent namespaces (implicit dotted hierarchy)",
            "Explicit lookups (published, read-only only)",
        ],
    )

    h1(doc, "Interfaces by role")
    table(
        doc,
        ["Role", "Primary tools"],
        [
            ["Developer", "MCP (Cursor/Claude), CLI, API"],
            ["QA / Tester", "Web UI, CLI (read), test plans tied to published facts"],
            ["Product owner", "Review queue, publish/approve, namespace design"],
            ["Business analyst", "Propose facts, document extraction, traceability"],
        ],
    )

    return doc


def build_facilitator_guide() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    cover(doc, "UniFact Training Program", "Facilitator Guide (Instructor Edition)")

    h1(doc, "Program overview")
    p(doc, "Duration: 3–4 days over 2 weeks (or 2 intensive days).")
    p(doc, "Audience: Developers, QA, Product owners, Business analysts.")
    p(doc, "Prerequisites: Training sandbox registry (see below), person keys issued, MCP configured for dev room.")
    bullets(
        doc,
        [
            "Foundation Handout — all roles (provided separately)",
            "One lab workbook per role",
            "Capstone requires mixed teams of 4",
        ],
    )

    h1(doc, "Training sandbox (facilitator setup)")
    p(
        doc,
        "The sandbox is a dedicated UniFact registry used only for training — not production org truth. "
        "Learners propose, publish, supersede, and break things safely without affecting real company facts.",
    )
    h2(doc, "What to create before Day 1")
    bullets(
        doc,
        [
            "A training registry (e.g. AcmeTraining) on your UniFact host or local store.db.",
            "Namespaces: product.demo, sales.policy, company.guidelines (copy or lookup to Unifact/company.guidelines).",
            "Person keys per attendee: dev-alice, qa-bob, po-carol, ba-dave (uni use <person>).",
            "Optional: lookup add from training registry to Unifact/company.guidelines for shared read-only basics.",
            "Sample policy PDF for BA extraction lab (can be a fake 'Acme Returns Policy').",
        ],
    )
    h2(doc, "Sandbox rules (tell the class)")
    bullets(
        doc,
        [
            "Proposed facts in the sandbox are for learning — not production policy.",
            "Never point MCP at production registry keys during labs unless explicitly doing a read-only demo.",
            "Reset option: new store.db or delete training registry facts between cohorts.",
            "Secrets (API keys, passwords) never go in facts — use env vars per company.infrastructure guidance.",
        ],
    )
    h2(doc, "Minimal local sandbox (facilitator machine)")
    bullets(
        doc,
        [
            "Clone unifact repo; npm install && npm run build.",
            "Use DATABASE_PATH=store.db (or a copy docs/training/sandbox-store.db).",
            "uni use <training-person>; seed or propose starter facts in product.demo.",
            "Point trainee .cursor/mcp.json at the same store.db for shared local labs, or per-laptop copies.",
        ],
    )

    h1(doc, "Day 1 — Foundation (all roles, 3.5 hours)")
    table(
        doc,
        ["Time", "Module", "Facilitator notes"],
        [
            ["0:00–0:15", "Welcome & objectives", "Emphasize: One Fact. One Truth. Separate docs vs managed facts."],
            ["0:15–1:00", "1.1 Why UniFact", "Use return-policy example. Ask: where is truth today?"],
            ["1:00–2:00", "1.2 Core concepts", "Use Foundation Handout. Quiz: document vs draft vs published."],
            ["2:00–2:15", "Break", ""],
            ["2:15–2:45", "1.3 Resolution order", "Whiteboard lookup path. Common mistake: lookup ≠ membership."],
            ["2:45–3:15", "1.4 Interfaces", "Demo MCP Fact Check + CLI list_facts."],
            ["3:15–3:45", "1.5 Live lifecycle demo", "Propose → approve → publish → supersede."],
        ],
    )

    h1(doc, "Day 1 PM / Day 2 — Role tracks")
    p(doc, "Split into breakouts. Each group uses their role lab workbook.")
    table(
        doc,
        ["Track", "Duration", "Workbook", "Certification"],
        [
            ["Developer", "6–7 hrs", "Lab-Developer", "MCP + propose fact + Fact Check"],
            ["QA", "4–5 hrs", "Lab-QA", "Lifecycle test plan + published-only assertion"],
            ["Product owner", "4–5 hrs", "Lab-ProductOwner", "Review queue + namespace map"],
            ["Business analyst", "4–5 hrs", "Lab-BusinessAnalyst", "10 proposed facts + traceability"],
        ],
    )

    h1(doc, "Capstone (2.5 hours, mixed teams)")
    h2(doc, "Scenario")
    p(doc, "Return policy changes from 30 to 45 days for Customer X (account-specific fact).")
    table(
        doc,
        ["Role", "Task", "Success criteria"],
        [
            ["BA", "Propose universal + account-specific facts with evidence", "Keys follow namespace convention"],
            ["PO", "Review, approve, publish; supersede old fact", "Only one published return-window per scope"],
            ["Dev", "Fact Check; reference published key in code/config", "No hardcoded 30 when published says 45"],
            ["QA", "Test published fact drives behavior", "Fails if proposed fact used in prod path"],
        ],
    )
    h2(doc, "Debrief questions (facilitator)")
    bullets(
        doc,
        [
            "What broke when someone used a proposed fact as truth?",
            "What's the difference between lookup and membership?",
            "When would you supersede vs edit in place?",
            "What belongs in company.infrastructure vs a published policy namespace?",
        ],
    )

    h1(doc, "Answer key — Foundation quiz")
    p(doc, "Document: 'Return policy v3.docx' → document, not a Uni Fact until extracted and published.")
    p(doc, "Proposed fact in review queue → draft, not truth.")
    p(doc, "Published sales.policy/returns → organizational truth for agents and apps.")

    h1(doc, "Troubleshooting (MCP / CLI)")
    table(
        doc,
        ["Symptom", "Fix"],
        [
            ["Cannot find dist/mcp.js", "npm run build in unifact repo; use absolute path"],
            ["better_sqlite3 NODE_MODULE_VERSION", "Rebuild native module or use system Node path"],
            ["Wrong propose attribution", "uni use <workAgentPerson> before agent writes"],
            ["Write denied on looked-up namespace", "Expected — lookup is read-only; join registry to write"],
        ],
    )

    h1(doc, "Materials checklist")
    bullets(
        doc,
        [
            "Sandbox registry (e.g. AcmeTraining) with training person keys",
            "Projector + facilitator machine with working MCP demo",
            "Sample policy PDF for BA extraction lab",
            "Printed or digital Foundation Handout for all attendees",
            "Certification sign-off sheet",
        ],
    )

    return doc


def build_dev_workbook() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    cover(doc, "UniFact Lab Workbook", "Developer Track")

    h1(doc, "Track outcomes")
    bullets(
        doc,
        [
            "Configure MCP and run Fact Check",
            "Use CLI: list, get, search, propose, sync",
            "Understand registry vs namespace and lookup (read-only)",
            "Integrate fact-dependent work into daily workflow",
        ],
    )

    h2(doc, "Module A1 — MCP setup (90 min)")
    exercise(
        doc,
        "A1.1",
        "Configure Cursor MCP",
        [
            "Copy .cursor/mcp.json.example to .cursor/mcp.json (or use global config).",
            "Set absolute path to node and dist/mcp.js; set DATABASE_PATH to store.db.",
            "Restart Cursor; verify user-unifact server appears in MCP tools.",
        ],
        3,
    )
    exercise(
        doc,
        "A1.2",
        "First Fact Check",
        [
            "In agent chat: run sync_pull (if upstream configured).",
            "search_facts for 'registry namespace' or list_facts on company.guidelines.",
            "Record one published fact you will rely on for today's work.",
        ],
        2,
    )

    h2(doc, "Module A2 — CLI (90 min)")
    exercise(
        doc,
        "A2.1",
        "Identity and list",
        [
            "uni use <your-training-person>",
            "uni list company.guidelines (or your sandbox namespace)",
            "uni get <namespace> <key> for one fact",
        ],
        3,
    )
    exercise(
        doc,
        "A2.2",
        "Propose a fact",
        [
            "Propose product.demo/training_complete = 'false' with description 'Developer lab flag'.",
            "Confirm it appears in proposed/working channel, not as published.",
            "Ask PO to approve in parallel track or use review queue if you have rights.",
        ],
        2,
    )

    h2(doc, "Module A3 — MCP tools (90 min)")
    exercise(
        doc,
        "A3.1",
        "Agent propose workflow",
        [
            "Give agent a task that depends on org policy (e.g. host redirect).",
            "Require Fact Check before edits.",
            "If fact missing: agent asks human, then propose_fact — does not assert in code.",
        ],
        3,
    )

    h2(doc, "Module A4 — Architecture (60 min)")
    exercise(
        doc,
        "A4.1",
        "Namespace map",
        [
            "Draw registries vs namespaces for your training sandbox.",
            "Mark one org-public namespace and one private namespace.",
            "Add one lookup path (read-only) to Unifact/company.guidelines.",
        ],
        4,
    )

    h2(doc, "Module A5 — API / CI (60 min)")
    exercise(
        doc,
        "A5.1",
        "Scripted fact check",
        [
            "Write a script or CI step that fails if a required published fact is missing.",
            "Document which namespace/key it guards.",
        ],
        2,
    )

    h1(doc, "Developer certification checklist")
    table(
        doc,
        ["Requirement", "Done?"],
        [
            ["MCP connected; sync_pull succeeds", "[ ]"],
            ["Proposed one fact with correct namespace/key", "[ ]"],
            ["Completed Fact Check on sample story", "[ ]"],
            ["Can explain lookup vs membership", "[ ]"],
        ],
    )

    return doc


def build_qa_workbook() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    cover(doc, "UniFact Lab Workbook", "QA / Tester Track")

    h1(doc, "Track outcomes")
    bullets(
        doc,
        [
            "Design tests around published vs proposed facts",
            "Validate full lifecycle in test environments",
            "Verify lookup is read-only; membership gates write",
            "Test agent behavior when facts are missing",
        ],
    )

    h2(doc, "Module B1 — Facts in test design (60 min)")
    exercise(
        doc,
        "B1.1",
        "Test case matrix",
        [
            "List 5 behaviors that depend on published facts (e.g. return window, feature flag).",
            "For each: namespace/key, expected published value, test env source.",
        ],
        5,
    )

    h2(doc, "Module B2 — Lifecycle testing (90 min)")
    exercise(
        doc,
        "B2.1",
        "Lifecycle path",
        [
            "With PO/Dev: propose a fact → verify NOT in prod behavior.",
            "After publish: verify prod/test uses new value.",
            "Supersede old fact: verify old value no longer authoritative.",
        ],
        4,
    )

    h2(doc, "Module B3 — Lookup & tenancy (60 min)")
    exercise(
        doc,
        "B3.1",
        "Negative tests",
        [
            "Confirm read via lookup works for org-public guidelines.",
            "Attempt write to looked-up namespace — expect failure.",
            "Document expected error or denial.",
        ],
        3,
    )

    h2(doc, "Module B4 — Agent testing (90 min)")
    exercise(
        doc,
        "B4.1",
        "Missing fact behavior",
        [
            "Prompt agent to implement feature requiring unpublished policy.",
            "Pass if agent asks/clarifies; fail if it invents policy in output.",
        ],
        3,
    )

    h2(doc, "Module B5 — Audit (60 min)")
    exercise(
        doc,
        "B5.1",
        "Audit trail",
        [
            "After a publish in sandbox, record: who approved, version, published_at.",
            "Write one regression case: 'policy change does not break unrelated facts'.",
        ],
        3,
    )

    h1(doc, "QA certification checklist")
    table(
        doc,
        ["Requirement", "Done?"],
        [
            ["Test plan covers lifecycle + published-only", "[ ]"],
            ["Negative test: lookup write denied", "[ ]"],
            ["Agent missing-fact test documented", "[ ]"],
            ["One bug/enhancement filed for draft-as-truth risk", "[ ]"],
        ],
    )

    return doc


def build_po_workbook() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    cover(doc, "UniFact Lab Workbook", "Product Owner Track")

    h1(doc, "Track outcomes")
    bullets(
        doc,
        [
            "Operate review queue: approve, reject, publish",
            "Design namespaces for a product area",
            "Decide what should become a Uni Fact",
            "Communicate truth changes to dev and QA",
        ],
    )

    h2(doc, "Module C1 — PO responsibilities (45 min)")
    exercise(
        doc,
        "C1.1",
        "Review queue",
        [
            "Open review queue (CLI or UI): list_review_queue / web equivalent.",
            "Review 5 proposed facts from BA/Dev sandbox.",
            "Approve 3, reject 2 with written rationale in workbook.",
        ],
        5,
    )

    h2(doc, "Module C2 — What is a fact? (60 min)")
    exercise(
        doc,
        "C2.1",
        "Story to fact",
        [
            "Take 3 user stories; extract candidate Uni Facts (namespace/key/value).",
            "Mark each: publish now / propose later / not a fact (stay in ticket).",
        ],
        4,
    )

    h2(doc, "Module C3 — Namespace design (60 min)")
    exercise(
        doc,
        "C3.1",
        "Namespace map",
        [
            "Design product.<yourapp>.* hierarchy for your team.",
            "Mark which namespaces are org-public vs private.",
            "Confirm no registry/namespace name collision.",
        ],
        6,
    )

    h2(doc, "Module C4 — Release & change (60 min)")
    exercise(
        doc,
        "C4.1",
        "Supersede drill",
        [
            "Publish v1 of a policy fact; announce to team.",
            "Publish v2 via supersede; document who must retest.",
        ],
        3,
    )

    h2(doc, "Module C5 — Service agents (optional, 45 min)")
    exercise(
        doc,
        "C5.1",
        "Facts for /build",
        [
            "List facts a customer service agent must have (hours, services, policies).",
            "Approve published set for demo agent.",
        ],
        3,
    )

    h1(doc, "Product owner certification checklist")
    table(
        doc,
        ["Requirement", "Done?"],
        [
            ["Approved/rejected facts with rationale", "[ ]"],
            ["Namespace map for one product area", "[ ]"],
            ["Led one supersede / publish communication", "[ ]"],
            ["Can explain org-public vs private", "[ ]"],
        ],
    )

    return doc


def build_ba_workbook() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    cover(doc, "UniFact Lab Workbook", "Business Analyst Track")

    h1(doc, "Track outcomes")
    bullets(
        doc,
        [
            "Write propose-ready Uni Facts with evidence",
            "Extract facts from documents",
            "Trace requirements to namespace/key",
            "Resolve conflicts via supersede workflow",
        ],
    )

    h2(doc, "Module D1 — Fact writing (60 min)")
    exercise(
        doc,
        "D1.1",
        "Ten facts",
        [
            "From sample policy PDF or provided doc, draft 10 facts.",
            "Each: namespace, key, value, evidence, suggested owner.",
        ],
        10,
    )

    h2(doc, "Module D2 — Extraction (90 min)")
    exercise(
        doc,
        "D2.1",
        "Document to propose",
        [
            "Run extract_facts_from_document (MCP or CLI) on sample doc.",
            "Curate output: merge duplicates, fix keys.",
            "Submit proposals; PO publishes subset in capstone.",
        ],
        4,
    )

    h2(doc, "Module D3 — Traceability (60 min)")
    exercise(
        doc,
        "D3.1",
        "Requirements matrix",
        [
            "Map 5 BRD requirements to Uni Fact keys.",
            "Add related_facts links where applicable.",
        ],
        5,
    )

    h2(doc, "Module D4 — Conflicts (60 min)")
    exercise(
        doc,
        "D4.1",
        "Conflicting returns",
        [
            "Given: '30-day return' and '45-day return' — propose resolution.",
            "Use account-specific vs universal namespaces; plan supersede.",
        ],
        3,
    )

    h2(doc, "Module D5 — Agile (45 min)")
    exercise(
        doc,
        "D5.1",
        "Acceptance criteria",
        [
            "Write one user story with AC referencing published fact key(s).",
            "Avoid duplicating policy text in Jira — link to Uni Fact.",
        ],
        3,
    )

    h1(doc, "Business analyst certification checklist")
    table(
        doc,
        ["Requirement", "Done?"],
        [
            ["10 well-formed proposed facts with evidence", "[ ]"],
            ["Traceability matrix complete", "[ ]"],
            ["Conflict resolution documented", "[ ]"],
            ["One story with AC → fact key references", "[ ]"],
        ],
    )

    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        ("UniFact-Training-Foundation-Handout.docx", build_foundation_handout()),
        ("UniFact-Training-Facilitator-Guide.docx", build_facilitator_guide()),
        ("UniFact-Training-Lab-Developer.docx", build_dev_workbook()),
        ("UniFact-Training-Lab-QA.docx", build_qa_workbook()),
        ("UniFact-Training-Lab-ProductOwner.docx", build_po_workbook()),
        ("UniFact-Training-Lab-BusinessAnalyst.docx", build_ba_workbook()),
    ]
    for name, document in files:
        path = OUT_DIR / name
        document.save(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
